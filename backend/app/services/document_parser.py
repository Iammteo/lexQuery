import io
from dataclasses import dataclass
from typing import List

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.models.document import DocumentType


@dataclass
class ParsedPage:
    """A single page of extracted text with metadata."""
    page_number: int
    text: str


@dataclass
class ParsedDocument:
    """
    Result of parsing a document.
    Pages are preserved separately so citations can reference
    specific page numbers later.
    """
    pages: List[ParsedPage]
    total_pages: int
    total_chars: int

    @property
    def full_text(self) -> str:
        """Concatenate all pages into a single text block."""
        return "\n\n".join(p.text for p in self.pages)


class DocumentParser:
    """
    Parse raw document bytes into structured text.
    Supports: PDF (with OCR fallback planned), DOCX.
    """

    @staticmethod
    def parse(file_bytes: bytes, document_type: DocumentType) -> ParsedDocument:
        """Parse a document and return structured text by page."""
        if document_type == DocumentType.PDF:
            return DocumentParser._parse_pdf(file_bytes)
        elif document_type == DocumentType.DOCX:
            return DocumentParser._parse_docx(file_bytes)
        elif document_type == DocumentType.TXT:
            return DocumentParser._parse_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> ParsedDocument:
        """
        Parse PDF using PyMuPDF.
        Each page is extracted separately so we can cite page numbers.
        """
        pages: List[ParsedPage] = []

        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            for page_num, page in enumerate(pdf, start=1):
                text = page.get_text("text").strip()
                if text:  # skip blank pages
                    pages.append(ParsedPage(page_number=page_num, text=text))
            total_pages = pdf.page_count
        finally:
            pdf.close()

        total_chars = sum(len(p.text) for p in pages)
        return ParsedDocument(
            pages=pages,
            total_pages=total_pages,
            total_chars=total_chars,
        )

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> ParsedDocument:
        """
        Parse DOCX using python-docx.
        DOCX has no real page concept — we treat the whole doc as one page.
        (For precise pagination we'd need LibreOffice/Word to render it.)
        """
        doc = DocxDocument(io.BytesIO(file_bytes))

        paragraphs: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also include tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n\n".join(paragraphs)

        return ParsedDocument(
            pages=[ParsedPage(page_number=1, text=full_text)],
            total_pages=1,
            total_chars=len(full_text),
        )

    @staticmethod
    def _parse_txt(file_bytes: bytes) -> ParsedDocument:
        """Parse plain text — no structure, just raw content."""
        text = file_bytes.decode("utf-8", errors="replace").strip()
        return ParsedDocument(
            pages=[ParsedPage(page_number=1, text=text)],
            total_pages=1,
            total_chars=len(text),
        )
